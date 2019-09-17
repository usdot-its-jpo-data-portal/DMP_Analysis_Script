import docx
import os
from glob import glob
from lxml import etree
from DDETable import DDETable
from DDERow import DDERow
from DDECell import DDECell
from DDEText import DDEText
from datetime import datetime


class DDEval:
    def __init__(self, params=None):
        self._path = None
        self._params = params

    def _get_chbox_value(self, sdtElement):
        val = ''
        ch_boxs = etree.ElementBase.xpath(sdtElement, './/w:sdtPr/w14:checkbox', namespaces=sdtElement.nsmap)
        if ch_boxs:
            for chbox in ch_boxs:
                ch = etree.ElementBase.xpath(chbox, './/w14:checked/@w14:val', namespaces=chbox.nsmap)
                if ch and len(ch) > 0:
                    val = ch[0]
            return val

    def _get_tc_value(self, tcElement):

        # Validates if the element belong to a grid
        try:
            tcElement._grid_col
        except:
            return None

        txt = ''
        tlst = etree.ElementBase.xpath(tcElement, './/w:p/w:r/w:t', namespaces=tcElement.nsmap)
        for t in tlst:
            txt += t.text

        return txt

    def _get_tables_information(self, doc):
        tables = []
        table_counter = 0
        for table in doc.tables:
            table_counter += 1

            dde_table = DDETable()
            dde_table.index = table_counter

            # print('{:_<80}'.format('_'))
            rows = table.rows
            row_counter = 0
            for row in rows:
                # print('{:_<50}'.format('_'))

                dde_row = DDERow()
                dde_row.index = row_counter

                row_counter += 1
                col_counter = 0

                # Select regular columns
                tcElements = etree.ElementBase.xpath(row._element, './/w:tc', namespaces=row._element.nsmap)
                if tcElements:
                    for tcElem in tcElements:
                        val = self._get_tc_value(tcElem)

                        if val is None:
                            continue


                        dde_cell = DDECell()
                        dde_cell.row_index = row_counter
                        dde_cell.column_index = col_counter
                        dde_cell.cell_type = 'text'
                        dde_cell.value = val

                        dde_row.add_cell(dde_cell)
                        # print('[{:d},{:d}] <{:s}>'.format(dde_cell.row_index, dde_cell.column_index, dde_cell.value))
                        col_counter += 1

                # Find checkboxes
                sdtElements = row._element.xpath('.//w:sdt')
                if sdtElements:
                    for sdt in sdtElements:
                        col_counter += 1
                        val = self._get_chbox_value(sdt)

                        dde_cell = DDECell()
                        dde_cell.row_index = row_counter
                        dde_cell.column_index = col_counter
                        dde_cell.cell_type = 'checkbox'
                        dde_cell.value = val

                        dde_row.add_cell(dde_cell)
                        # print('[{:d},{:d}] <{:s}>'.format(dde_cell.row_index, dde_cell.column_index, dde_cell.value))

                dde_table.add_row(dde_row)

            dde_table.name = dde_table.row(0).cell(1).value
            tables.append(dde_table)

        return tables

    def _get_paragraphs_information(self, doc):
        result = []
        for paragraph in doc.paragraphs:

            sdtContents = etree.ElementBase.xpath(paragraph._element, './/w:sdt/w:sdtContent', namespaces=paragraph._element.nsmap)
            if sdtContents:
                txt = paragraph.text
                t_lst = etree.ElementBase.xpath(sdtContents[0], './/w:r/w:t', namespaces=sdtContents[0].nsmap)
                val = ''
                if t_lst:
                    for t in t_lst:
                        val += t.text

                dde_text = DDEText()
                dde_text.text = txt
                dde_text.value = val if len(val) > 1 else '1' if ord(val) == 9746 else '0' if ord(val) == 9744 else val

                dde_text.resolve_id()

                result.append(dde_text)

        return result

    def _generate_detailed_report(self, path_name, text_fields, tables,  separator='\t'):
        lines = []
        rep_meta = os.path.basename(path_name)

        for text_field in text_fields:
            rep_meta += '{}{}{}{}{}{}'.format(separator, text_field.id, separator, text_field.text, separator, text_field.value)

        for table in tables:
            for row in table.rows:
                rep_line = rep_meta + '{}{}'.format(separator, table.name)
                for cell in row.cells:
                    rep_line += '{}{}'.format(separator, cell.value)

                lines.append('{}'.format(rep_line))

        return lines

    def _generate_summary_report(self, tables):
        dic_summary = {}
        lst_vals_title = []
        lst_vals_counters = []
        for table in tables:
            if 'ID' not in dic_summary:
                if table.row_count() > 0:  # if the dic_summary is empty
                    r = table.row(0)
                    i = 2  # index to the first checkbox
                    while i < r.count():
                        c = r.cell(i)
                        if c.value not in lst_vals_title:
                            lst_vals_title.append(c.value)
                            lst_vals_counters.append(0)

                        i += 1
                    lst_vals_title.insert(0, 'Description')
                    lst_vals_counters.insert(0, '')
                    dic_summary['ID'] = lst_vals_title

            for row in table.rows:

                # Get Row ID
                cell = row.cell(0)
                if cell is None:
                    continue
                val = cell.value
                if not val:
                    continue

                # Get Row description
                val_desc = ''
                cell_desc = row.cell(1)
                if cell_desc is not None:
                    val_desc = cell_desc.value

                vals = [x for x in lst_vals_counters]
                vals[0] = val_desc
                if val in dic_summary:
                    vals = dic_summary[val]

                i = 1
                for c in row.cells:
                    if c.cell_type == 'checkbox':
                        vals[i] = vals[i] + int(c.value)
                        i += 1

                dic_summary[val] = vals

        return dic_summary

    def _save_report(self, data, num_files, separator='\t'):

        target_filename = None
        if self._params.output is not None:
            if os.path.isdir(self._params.output):
                id_str = 'detailed' if self._params.detailed else 'summary'
                if self._params.is_directory:
                    dt_str = datetime.now().strftime('%Y%m%d%H%M%S')
                    target_filename = 'DMP_dir_{}_[{}].csv'.format(id_str, dt_str);
                else:
                    fn = os.path.basename(self._params.path_name)
                    target_filename = '{}_{}{}'.format(fn, id_str, '.csv')


                target_filename = os.path.join(self._params.output, target_filename)

            else:
                print('Error: Invalid output directory')
                return

        if self._params.detailed:
            if type(data) is list:
                if target_filename is not None:
                    with open(target_filename, 'w') as file:
                        for line in data:
                            file.write('{}\n'.format(line))
                    print('Report saved at: {}'.format(target_filename))
                else:
                    for line in data:
                        print(line)
        else:
            if type(data) is dict:
                if target_filename is not None:
                    with open(target_filename, 'w') as file:
                        file.write('{}{}{}\n'.format('Number of files: ', separator, num_files))
                        for key in data:
                            vals = data[key]
                            rep_line = key
                            for val in vals:
                                rep_line += '{}{}'.format(separator, val)
                            file.write('{}\n'.format(rep_line))
                    print('Report saved at: {}'.format(target_filename))
                else:
                    print('{}{}{}\n'.format('Number of files: ', separator, num_files))
                    for key in data:
                        vals = data[key]
                        rep_line = key
                        for val in vals:
                            rep_line += '{}{}'.format(separator, val)
                        print(rep_line)

    def _process_file(self, filename):
        doc = docx.Document(filename);

        text_fields = self._get_paragraphs_information(doc)
        tables = self._get_tables_information(doc)

        if self._params.detailed:
            lines = self._generate_detailed_report(filename, text_fields, tables)
            return lines
        else:
            dic_summary = self._generate_summary_report(tables)
            return dic_summary

    def _process_directory(self):
        files = glob(os.path.join(self._params.path_name, '*.docx'))
        if len(files) <= 0:
            print('Error: No files found.')
            return

        lines = []
        dic_summary = {}
        for filename in files:
            file_result = self._process_file(filename)
            if type(file_result) is dict:
                for key in file_result:
                    if key in dic_summary:
                        if key == 'ID':
                            continue
                        else:
                            file_data = file_result[key]
                            sum_data = dic_summary[key]
                            i = 1
                            while i < len(file_data):
                                sum_data[i] = sum_data[i] + file_data[i]
                                i += 1

                            dic_summary[key] = sum_data
                    else:
                        dic_summary[key] = file_result[key]

            else:
                if type(file_result) is list:
                    for line in file_result:
                        lines.append(line)

        if self._params.detailed:
            return lines, len(files)
        else:
            return dic_summary, len(files)

    def parse(self):
        if self._params is None:
            print('Error: Invalid parameters')
            exit(1)

        if self._params.is_directory:
            result, num_files = self._process_directory()
            self._save_report(result, num_files)

        else:
            result = self._process_file(self._params.path_name)
            self._save_report(result, 1)





